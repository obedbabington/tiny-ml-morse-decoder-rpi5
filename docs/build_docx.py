#!/usr/bin/env python3
"""
Build the publication-ready Word document for the MorseAI Edge AI tutorial.

Generates docs/MorseAI_EdgeAI_RaspberryPi5.docx from python-docx, including a
title page, executive summary, figures with captions, and result tables.

Run (from repo root or docs/):
    python3 docs/build_docx.py

Requires python-docx. The script prefers the vendored copy in .build_tools/ if
present; otherwise install with:
    pip3 install python-docx
"""

import sys
from pathlib import Path

# Prefer the project's vendored python-docx over any broken global "docx" package.
_REPO_ROOT = Path(__file__).resolve().parent.parent
_BUILD_TOOLS = _REPO_ROOT / ".build_tools"
if _BUILD_TOOLS.is_dir():
    sys.path.insert(0, str(_BUILD_TOOLS))

try:
    from docx import Document
    import docx as _docx
except ModuleNotFoundError as exc:
    raise SystemExit(
        "Could not import python-docx.\n\n"
        "Install it with:\n"
        "  pip3 install python-docx\n\n"
        "Or, from the repo root, vend a local copy:\n"
        "  pip3 install --target=.build_tools python-docx\n"
    ) from exc

# Guard against the obsolete PyPI package named "docx" (a single docx.py file).
_docx_file = getattr(_docx, "__file__", "") or ""
if _docx_file.endswith("docx.py") or not hasattr(_docx, "shared"):
    raise SystemExit(
        "Wrong 'docx' package is installed (not python-docx).\n\n"
        "Fix:\n"
        "  pip3 uninstall docx\n"
        "  pip3 install python-docx\n\n"
        "Or use the vendored copy:\n"
        "  pip3 install --target=.build_tools python-docx\n"
        "  python3 docs/build_docx.py\n"
    )
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

HERE = Path(__file__).parent
FIG = HERE / "architecture.png"
OUT = HERE / "MorseAI_EdgeAI_RaspberryPi5.docx"

ACCENT = RGBColor(0x2B, 0x6C, 0xB0)
INK = RGBColor(0x1F, 0x29, 0x33)
MUTED = RGBColor(0x52, 0x60, 0x6D)


def set_base_styles(doc: Document) -> None:
    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15

    for name, size, color in [
        ("Heading 1", 16, ACCENT),
        ("Heading 2", 13, INK),
        ("Heading 3", 11.5, INK),
    ]:
        st = doc.styles[name]
        st.font.name = "Calibri"
        st.font.size = Pt(size)
        st.font.color.rgb = color
        st.font.bold = True


def caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.italic = True
    run.font.size = Pt(9.5)
    run.font.color.rgb = MUTED


def shade_header_row(table) -> None:
    for cell in table.rows[0].cells:
        tcPr = cell._tc.get_or_add_tcPr()
        shd = tcPr.makeelement(qn("w:shd"), {
            qn("w:val"): "clear",
            qn("w:color"): "auto",
            qn("w:fill"): "2B6CB0",
        })
        tcPr.append(shd)
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)


def add_table(doc: Document, headers, rows, widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.paragraphs[0].add_run(h).bold = True
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
    if widths:
        for col, w in enumerate(widths):
            for r in table.rows:
                r.cells[col].width = Inches(w)
    shade_header_row(table)
    doc.add_paragraph()
    return table


def bullets(doc: Document, items):
    for it in items:
        if isinstance(it, tuple):
            p = doc.add_paragraph(style="List Bullet")
            p.add_run(it[0]).bold = True
            p.add_run(" " + it[1])
        else:
            doc.add_paragraph(it, style="List Bullet")


def body(doc: Document, text: str) -> None:
    doc.add_paragraph(text)


def build() -> None:
    doc = Document()
    set_base_styles(doc)

    # ----------------------------------------------------------------- Title page
    for _ in range(3):
        doc.add_paragraph()
    t = doc.add_paragraph()
    t.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = t.add_run("MorseAI")
    r.font.size = Pt(34)
    r.font.bold = True
    r.font.color.rgb = ACCENT

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("An End-to-End Edge AI Workflow on Raspberry Pi 5")
    r.font.size = Pt(16)
    r.font.color.rgb = INK

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run(
        "A beginner-friendly tutorial for building and deploying neural "
        "networks on Arm-based edge hardware"
    )
    r.font.size = Pt(12)
    r.italic = True
    r.font.color.rgb = MUTED

    for _ in range(2):
        doc.add_paragraph()

    if FIG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIG), width=Inches(6.0))

    for _ in range(2):
        doc.add_paragraph()

    meta = [
        ("Project", "MorseAI: Raspberry Pi 5 port of an Arm STM32 workshop project"),
        ("Platform", "Raspberry Pi 5 (Arm Cortex-A76), 64-bit Raspberry Pi OS"),
        ("Framework", "TensorFlow / TensorFlow Lite (tflite-runtime)"),
        ("Author", "Obed Babington"),
        ("Status", "Tutorial documentation"),
        ("Date", "______________________"),
    ]
    mt = doc.add_table(rows=0, cols=2)
    mt.alignment = WD_TABLE_ALIGNMENT.CENTER
    for k, v in meta:
        cells = mt.add_row().cells
        rk = cells[0].paragraphs[0].add_run(k)
        rk.bold = True
        rk.font.color.rgb = INK
        cells[1].paragraphs[0].add_run(v)
        cells[0].width = Inches(1.6)
        cells[1].width = Inches(4.4)

    doc.add_page_break()

    # ----------------------------------------------------------- Executive summary
    doc.add_heading("Executive Summary", level=1)
    body(
        doc,
        "MorseAI is a compact, complete example of an Edge AI application: a user "
        "taps a letter in Morse code on a single push button, and a neural network "
        "running entirely on a Raspberry Pi 5 identifies the letter in real time. "
        "There is no cloud service and no network round-trip: capture, "
        "preprocessing, and inference all happen on the device.",
    )
    body(
        doc,
        "The project began as an Arm workshop on an STM32 microcontroller and was "
        "ported to the Raspberry Pi 5 to serve as an accessible teaching platform. "
        "Its value is not the Morse problem itself but the workflow it demonstrates: "
        "capturing a sensor signal, turning it into features, training a small model "
        "off-device, converting it to TensorFlow Lite, and running inference on Arm "
        "hardware. That same four-stage pattern underpins practical edge workloads "
        "such as sensor classification, vibration-based predictive maintenance, "
        "gesture recognition, and anomaly detection.",
    )
    body(
        doc,
        "The deployed model is approximately 4.2 KB and imposes negligible load on "
        "the Pi 5, leaving ample headroom to scale up to real sensors and larger "
        "models. This document describes the motivation, architecture, "
        "implementation decisions, measured resource cost, and the paths from this "
        "tutorial to production Edge AI systems, including a comparison with the "
        "original microcontroller implementation.",
    )

    # ------------------------------------------------------------------- 1. Intro
    doc.add_heading("1. Introduction and Motivation", level=1)
    body(
        doc,
        "Most introductions to machine learning stop at a notebook: a model is "
        "trained, an accuracy number is reported, and the work ends there. The "
        "harder and more useful skill is getting a model off the laptop and onto "
        "hardware where it has to deal with real, messy input. MorseAI was built to "
        "teach exactly that transition, using the smallest honest example we could "
        "find.",
    )
    body(
        doc,
        "Morse code is a good first signal to classify because each character is "
        "just a handful of button presses that differ only in duration. The data is "
        "tiny and human-readable, you can generate your own dataset in minutes, and "
        "the problem is genuinely temporal: the same class of concern (timing, "
        "noise, sampling) that appears in audio and vibration, but slow enough to "
        "debug by hand. It is explicitly a learning example, not an industrial "
        "product, and it is framed that way throughout.",
    )

    # -------------------------------------------------------------- 2. Background
    doc.add_heading("2. Why Morse Code Is a Good First Edge AI Problem", level=1)
    bullets(doc, [
        ("Small, interpretable input.", "Each character is at most four press "
         "durations: four numbers you can read and sanity-check by eye."),
        ("Genuinely temporal.", "A dot and a dash differ only in duration, which "
         "forces real thinking about timing and noise."),
        ("You own the data.", "A usable dataset is collected by pressing a button, "
         "so the learner controls the whole pipeline end to end."),
        ("Instructive failure modes.", "Ambiguous taps show why a learned model is "
         "more forgiving than a hand-tuned threshold, motivating the use of ML "
         "in the first place."),
    ])

    # ------------------------------------------------------------ 3. Architecture
    doc.add_heading("3. System Architecture", level=1)
    body(
        doc,
        "The system is a four-stage pipeline. Each stage maps to a single source "
        "file, which keeps every part independently testable. The training step "
        "runs once, off-device; everything in Figure 1 runs on the Pi.",
    )
    if FIG.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(FIG), width=Inches(6.2))
    caption(doc, "Figure 1. The MorseAI Edge AI pipeline: signal capture \u2192 "
                 "preprocessing \u2192 TensorFlow Lite inference \u2192 decoded output, "
                 "running on-device on a Raspberry Pi 5.")
    add_table(
        doc,
        ["Stage", "File", "Responsibility"],
        [
            ["Signal capture", "capture.py", "Time button presses, debounce, detect end-of-character"],
            ["Preprocessing", "inference.py", "Pad to four values, apply the training StandardScaler"],
            ["Inference", "inference.py", "Run the TFLite interpreter; return top class + confidence"],
            ["Output", "main.py", "Wire capture to inference; print results; hold session state"],
        ],
        widths=[1.4, 1.4, 3.6],
    )
    caption(doc, "Table 1. Mapping of pipeline stages to source files.")

    # --------------------------------------------------------- 4. Implementation
    doc.add_heading("4. Implementation", level=1)

    doc.add_heading("4.1 Signal Capture", level=2)
    body(
        doc,
        "The button is read with the gpiozero library, which fires callbacks on "
        "press and release. A monotonic, high-resolution clock (perf_counter_ns) "
        "measures each press duration in microseconds. Two implementation details "
        "matter more than they appear: a 50 ms debounce filter, so mechanical "
        "contact chatter registers as a single press; and a 2-second "
        "end-of-character timeout, implemented as a small cancel-and-rearm timer "
        "that decides when a character is complete. Each character is stored as up "
        "to four durations and zero-padded to a fixed length of four.",
    )

    doc.add_heading("4.2 Preprocessing and Feature Representation", level=2)
    body(
        doc,
        "Raw press durations span tens to hundreds of thousands of microseconds. A "
        "StandardScaler (subtract mean, divide by standard deviation) brings them "
        "into a range the network trains on efficiently. Critically, the identical "
        "transform must be applied at inference time, so the scaler parameters are "
        "exported with the model and reloaded on the device. A normalization "
        "mismatch is the most common cause of a model that loads correctly but "
        "predicts nonsense.",
    )

    doc.add_heading("4.3 Model Design and Training", level=2)
    body(
        doc,
        "The classifier is a small feed-forward network: two hidden layers of 16 "
        "ReLU units feeding an 11-way softmax (ten letters A\u2013J plus one "
        "\u201Cunclassified\u201D class). With only four inputs and ten classes, depth "
        "is unnecessary and a larger model would simply overfit the 450-sample "
        "dataset. The eleventh class is trained on a small fraction of random noise "
        "vectors so the model can respond \u201CI don\u2019t know\u201D rather than forcing "
        "garbage onto the nearest letter, a small change that makes the deployed "
        "system noticeably more robust. Training uses Adam, sparse categorical "
        "cross-entropy, and early stopping on validation loss.",
    )

    doc.add_heading("4.4 Conversion to TensorFlow Lite", level=2)
    body(
        doc,
        "The trained Keras model is converted with the TFLite converter using "
        "default optimizations (weight quantization). Conversion produces two "
        "artifacts: the model file (4,276 bytes) and a normalization configuration "
        "file holding the scaler mean and scale. Shipping the normalization "
        "parameters alongside the model is what makes the result reproducible on a "
        "different machine.",
    )

    doc.add_heading("4.5 On-Device Inference", level=2)
    body(
        doc,
        "Inference on the Pi uses tflite-runtime rather than full TensorFlow: it is "
        "a few megabytes instead of hundreds, starts quickly, and is sufficient for "
        "running (not training) a model. A prediction normalizes the input, sets "
        "the input tensor, invokes the interpreter, and takes the argmax of the "
        "softmax output, returning the predicted letter with a confidence score and "
        "a measured inference time.",
    )

    # ------------------------------------------------------------- 5. Performance
    doc.add_heading("5. Performance on Raspberry Pi 5", level=1)
    body(
        doc,
        "Edge AI is judged on resource cost. The figures in Table 2 were measured "
        "on a Raspberry Pi 5 with python benchmark.py (10,000 iterations, 4 threads, "
        "200 warmup). TFLite loads the XNNPACK CPU delegate automatically.",
    )
    add_table(
        doc,
        ["Metric", "Value", "How measured"],
        [
            ["Model size (disk)", "4.18 KB (4,276 bytes)", "ls -l of the .tflite file"],
            ["Input / output", "4 \u00d7 float32 \u2192 11 \u00d7 float32", "model signature"],
            ["Inference latency (median)", "3.0 \u00b5s", "warm invoke() calls"],
            ["Inference latency (mean / p95)", "3.1 / 3.1 \u00b5s", "same run"],
            ["Inference latency (min / max)", "3.0 / 79.8 \u00b5s", "same run"],
            ["Throughput", "~280,400 inferences/s", "tight loop, 4 threads"],
            ["Interpreter memory (RSS \u0394)", "2.7 MiB", "39.5 MiB before \u2192 42.2 MiB after load"],
            ["Process RSS after load", "42.2 MiB", "/proc resident set size"],
            ["CPU during benchmark loop", "100%", "process CPU time \u00f7 wall time"],
        ],
        widths=[2.0, 2.0, 3.0],
    )
    caption(doc, "Table 2. Measured resource cost of the MorseAI model on Raspberry Pi 5.")
    body(
        doc,
        "Latency is far below one millisecond. In normal use with main.py, the "
        "bottleneck is the 2-second character timeout, not inference. A 4 KB model "
        "leaves large headroom for real sensors and bigger networks.",
    )
    body(
        doc,
        "CPU note: the benchmark runs a tight loop with no idle time between calls. "
        "100% means the process kept one core busy for the timed section. It does not "
        "mean all four Pi cores were saturated, and it is not typical of CPU use while "
        "waiting for button input.",
    )
    body(
        doc,
        "To reproduce: activate the project venv and run python benchmark.py from "
        "the repo root.",
    )

    # --------------------------------------------------------------- 6. Extending
    doc.add_heading("6. Extending the Pattern to Real-World Workloads", level=1)
    body(
        doc,
        "The pipeline does not change when the problem becomes serious. Replace the "
        "button with a different sensor, retrain on the new data, and redeploy the "
        "same way: capture \u2192 preprocess \u2192 TFLite inference \u2192 act. The "
        "following applications reuse the same skeleton.",
    )
    add_table(
        doc,
        ["Application", "What changes", "What stays the same"],
        [
            ["Sensor classification", "Accelerometer / environmental sensor instead of a button", "Fixed-length feature vector, small MLP"],
            ["Predictive maintenance / vibration", "Windowed vibration signal; normal vs. fault label", "Time-series in, classification out"],
            ["Gesture recognition", "IMU data from a wearable", "Variable-length input \u2192 fixed feature trick"],
            ["Anomaly detection", "Learn \u201Cnormal\u201D; flag drift", "The \u201Cunclassified\u201D idea, generalised"],
            ["Other IoT edge workloads", "Keyword spotting, occupancy, fault detection", "On-device decision, no cloud round-trip"],
        ],
        widths=[1.8, 2.6, 2.6],
    )
    caption(doc, "Table 3. The same capture-to-inference pipeline applied to "
                 "practical Edge AI workloads.")
    body(
        doc,
        "The step up from this tutorial is usually more and better data, a slightly "
        "larger model, and windowing a continuous stream rather than discrete "
        "presses; none of which change the deployment story.",
    )

    # ---------------------------------------------------------------- 7. Compare
    doc.add_heading("7. STM32 vs. Raspberry Pi 5: Two Faces of Edge AI", level=1)
    body(
        doc,
        "Because the project exists on both an STM32 microcontroller and a "
        "Raspberry Pi 5, it offers a clear side-by-side of two points on the Edge "
        "AI spectrum. Neither is strictly better; they answer different questions.",
    )
    add_table(
        doc,
        ["Dimension", "STM32 (microcontroller)", "Raspberry Pi 5 (Linux SBC)"],
        [
            ["Developer experience", "Bare-metal / RTOS, C/C++, flash-and-debug", "Full Linux, Python, edit-and-run"],
            ["Toolchain", "TFLite Micro, C compiler, model as C array", "TensorFlow \u2192 TFLite, tflite-runtime, pip/venv"],
            ["Deployment workflow", "Cross-compile and flash firmware", "git pull and run a script"],
            ["Compute & memory", "Hundreds of KB RAM, MHz core, no OS", "GBs of RAM, GHz quad Cortex-A76, full OS"],
            ["Iteration speed", "Slow: recompile + reflash", "Fast: edit a line, rerun in seconds"],
            ["Power & cost", "Milliwatts; cents to dollars", "Watts; tens of dollars"],
            ["Best for", "Always-on, battery, cost-sensitive, real-time", "Prototyping, heavier models, learning, gateways"],
        ],
        widths=[1.6, 2.7, 2.7],
    )
    caption(doc, "Table 4. Developer-facing comparison of the STM32 and "
                 "Raspberry Pi 5 implementations.")
    body(
        doc,
        "A realistic workflow uses both: develop and validate the model on the Pi, "
        "where the fast edit-run loop keeps attention on the problem, then port the "
        "proven model down to the STM32 with TFLite Micro when microwatt power, "
        "cents-level unit cost, or hard real-time guarantees are required. The Pi is "
        "the workshop; the microcontroller is the finished product.",
    )

    # --------------------------------------------------------------- 8. Lessons
    doc.add_heading("8. Challenges and Lessons Learned", level=1)
    bullets(doc, [
        ("The signal is harder than the network.", "Debounce, the end-of-character "
         "timeout, and padding took more effort than the model, typical for Edge AI."),
        ("Normalization is a deployment concern.", "Shipping the scaler with the "
         "model is essential; a mismatch yields confident nonsense."),
        ("An explicit \u201CI don\u2019t know\u201D class pays off.", "It lets the system "
         "degrade gracefully on ambiguous input."),
        ("tflite-runtime vs. TensorFlow is a real choice.", "The runtime is the "
         "right tool for inference-only deployment and shaped the Python setup."),
        ("Threads don\u2019t always help tiny models.", "Per-call overhead can "
         "dominate; benchmark.py makes it easy to compare 1, 2, and 4 threads."),
    ])

    # ------------------------------------------------------------- 9. Conclusion
    doc.add_heading("9. Conclusion and Future Work", level=1)
    body(
        doc,
        "MorseAI shows that a complete Edge AI workflow, from sensor capture to "
        "on-device inference, fits in a few small, readable files on a Raspberry "
        "Pi 5. The Morse problem is intentionally simple so that the workflow, which "
        "transfers directly to practical sensor and IoT applications, stays in "
        "focus. For a beginner, it is a realistic first deployment; for a "
        "practitioner, it is a clean template to adapt.",
    )
    bullets(doc, [
        "Grow and balance the dataset; extend coverage to full A\u2013Z and digits.",
        "Decode whole words using inter-character and inter-word gaps.",
        "Add an int8-quantized variant and compare accuracy against size.",
        "Provide TFLite Micro build notes for porting back to the STM32 target.",
    ])

    # ----------------------------------------------------------------- Appendices
    doc.add_heading("Appendix A: Hardware Setup and Wiring", level=1)
    add_table(
        doc,
        ["Button terminal", "Pi physical pin", "Pi function", "Role"],
        [
            ["Terminal 1", "Pin 11", "GPIO 17 (BCM)", "Signal in"],
            ["Terminal 2", "Pin 9", "GND", "Ground"],
        ],
        widths=[1.7, 1.6, 1.7, 1.5],
    )
    caption(doc, "Table A1. Button wiring. The code enables the internal pull-up, "
                 "so no external resistor is required.")

    doc.add_heading("Appendix B: Reproducing the Results", level=1)
    bullets(doc, [
        "Train and convert the model with training/MorseAI_RPi_Notebook.ipynb.",
        "Deploy on the Pi with setup_pi.sh, then run python main.py.",
        "Benchmark inference with python benchmark.py (no GPIO required).",
        "Regenerate the architecture figure with docs/make_diagram.py.",
    ])

    doc.save(str(OUT))
    print(f"Saved {OUT}")


if __name__ == "__main__":
    build()
